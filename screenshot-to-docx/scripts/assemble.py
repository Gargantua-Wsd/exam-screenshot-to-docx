from __future__ import annotations

import argparse
import json
import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageEnhance, ImageFilter, ImageOps
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm


def load_input(path: Path) -> Image.Image:
    files = [path] if path.is_file() else sorted(
        p for p in path.iterdir() if p.suffix.lower() in {".jpg", ".jpeg", ".png", ".webp"}
    )
    if not files:
        raise FileNotFoundError(f"No supported images found: {path}")
    images = [Image.open(p).convert("RGB") for p in files]
    width = max(i.width for i in images)
    total = sum(i.height for i in images)
    canvas = Image.new("RGB", (width, total), "white")
    y = 0
    for image in images:
        canvas.paste(image, ((width - image.width) // 2, y))
        y += image.height
    return canvas


def row_ink(image: Image.Image) -> list[float]:
    gray = ImageOps.grayscale(image)
    # Downsample for stable whitespace analysis; dark pixels are likely content.
    scale = max(1, image.width // 900)
    small = gray.resize((max(1, image.width // scale), max(1, image.height // scale)))
    values = []
    for y in range(small.height):
        row = [p for p in small.crop((0, y, small.width, y + 1)).getdata()]
        values.append(sum(p < 220 for p in row) / max(1, len(row)))
    return values


def auto_cuts(image: Image.Image) -> tuple[list[int], list[str]]:
    ink = row_ink(image)
    # Long blank bands are candidate separators. Merge nearby candidates.
    blank = max(4, image.height // 1200)
    candidates = []
    start = None
    for y, density in enumerate(ink):
        if density < 0.008:
            start = y if start is None else start
        elif start is not None:
            if y - start >= blank:
                candidates.append((start + y) // 2)
            start = None
    if start is not None and len(ink) - start >= blank:
        candidates.append((start + len(ink)) // 2)
    # Convert low-resolution coordinates back to source coordinates.
    scale = image.height / max(1, len(ink))
    candidates = [int(c * scale) for c in candidates]
    # Only use separators that create reasonably sized segments. The screenshot
    # may contain many small blank bands inside a question, so first collect
    # coarse pieces and then merge them into question-sized blocks.
    cuts = [0]
    low_confidence = []
    min_segment = max(220, image.height // 80)
    for c in candidates:
        if c - cuts[-1] >= min_segment:
            cuts.append(c)
    if image.height - cuts[-1] >= 80:
        cuts.append(image.height)
    else:
        cuts[-1] = image.height
    # Mobile exam pages often have several blank bands inside one question
    # (between stem, diagram, options, and feedback controls). Merge short
    # pieces until a plausible question block is reached. This is intentionally
    # conservative: under-splitting is easier to detect and manually correct
    # than silently losing part of a question.
    merged = [cuts[0]]
    target = max(720, min(1200, image.height // 24))
    for c in cuts[1:-1]:
        if c - merged[-1] >= target:
            merged.append(c)
    merged.append(cuts[-1])
    cuts = merged
    # Very tall segments are not necessarily wrong; flag them for human review.
    for a, b in zip(cuts, cuts[1:]):
        if b - a > image.height * 0.22:
            low_confidence.append(f"segment {a}:{b} is unusually tall")
    return cuts, low_confidence


def crop_questions(image: Image.Image, cuts: list[int], out: Path) -> list[Path]:
    out.mkdir(parents=True, exist_ok=True)
    result = []
    for index, (top, bottom) in enumerate(zip(cuts, cuts[1:]), 1):
        # Keep generous vertical context so question numbers and conditions do
        # not disappear when the separator falls inside a website card. The
        # final portion of a card is reserved for the site's post-question
        # controls; remove only that conservative tail band.
        top_pad = max(80, image.width // 12)
        tail_pad = max(80, image.width // 14)
        crop_top = max(0, top - top_pad)
        crop_bottom = image.height if bottom == image.height else max(crop_top + 100, bottom - tail_pad)
        crop = image.crop((0, crop_top, image.width, crop_bottom))
        crop = ImageEnhance.Contrast(crop).enhance(1.05).filter(ImageFilter.SHARPEN)
        filename = out / f"question_{index:03d}.png"
        crop.save(filename, optimize=True)
        result.append(filename)
    return result


def make_preview(image: Image.Image, cuts: list[int], out: Path) -> None:
    preview = image.copy()
    preview.thumbnail((1000, 1600))
    scale = preview.height / image.height
    draw = ImageDraw.Draw(preview)
    for index, y in enumerate(cuts[1:-1], 1):
        py = int(y * scale)
        draw.line((0, py, preview.width, py), fill=(220, 30, 30), width=3)
        draw.text((8, max(0, py - 20)), f"Q{index + 1}", fill=(220, 30, 30))
    preview.save(out)


def add_answer_page(doc: Document, style: str, number: int) -> None:
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"绗?{number} 棰樼瓟棰樺尯鍩?)
    run.bold = True
    if style == "blank":
        return
    if style == "grid":
        for _ in range(18):
            doc.add_paragraph("鈻?" * 24)
    else:
        for _ in range(18):
            doc.add_paragraph("________________________________________________________________")


def build_docx(question_files: list[Path], output: Path, answer_style: str) -> dict:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    usable_mm = 297 - 36
    long_questions = []
    answer_area_sizes = []
    for index, file in enumerate(question_files, 1):
        with Image.open(file) as image:
            ratio = image.height / max(1, image.width)
            height_mm = 210 * ratio
        is_long = height_mm >= usable_mm * 0.82
        if index > 1:
            doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(file), width=Mm(174))
        # Every question gets response space. Keep short-question space in the
        # current flow when possible, but give long questions a fresh full page.
        if is_long:
            long_questions.append(index)
            answer_area_sizes.append("full_page")
            add_answer_page(doc, answer_style, index)
        else:
            answer_area_sizes.append("half_page")
            add_answer_area(doc, answer_style, index, lines=9)
    doc.save(output)
    return {"question_count": len(question_files), "long_questions": long_questions, "answer_area_count": len(answer_area_sizes), "answer_area_sizes": answer_area_sizes, "answer_page_count": len(long_questions)}


def add_answer_page(doc: Document, style: str, number: int) -> None:
    doc.add_page_break()
    for _ in range(18):
        doc.add_paragraph()


def add_answer_area(doc: Document, style: str, number: int, lines: int) -> None:
    for _ in range(lines):
        doc.add_paragraph()


def main() -> None:
    parser = argparse.ArgumentParser(description="Split exam screenshots and assemble a DOCX")
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--split-mode", choices=["auto", "manual"], default="auto")
    parser.add_argument("--cuts", default="", help="Comma-separated pixel boundaries including optional 0/end")
    parser.add_argument("--answer-style", choices=["ruled", "blank", "grid"], default="ruled")
    args = parser.parse_args()
    args.output.mkdir(parents=True, exist_ok=True)
    image = load_input(args.input)
    if args.split_mode == "manual":
        cuts = [int(v) for v in args.cuts.split(",") if v.strip()]
        if not cuts or cuts[0] != 0:
            cuts.insert(0, 0)
        if cuts[-1] != image.height:
            cuts.append(image.height)
        warnings = []
    else:
        cuts, warnings = auto_cuts(image)
    question_dir = args.output / "question_images"
    questions = crop_questions(image, cuts, question_dir)
    make_preview(image, cuts, args.output / "question_boundaries.png")
    docx_info = build_docx(questions, args.output / "exam_from_screenshot.docx", args.answer_style)
    report = {"input": str(args.input), "source_size": [image.width, image.height], "cuts": cuts, "warnings": warnings, **docx_info}
    (args.output / "processing_report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()

